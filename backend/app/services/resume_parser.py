import io
import re
import json
from pathlib import Path
from typing import Optional

from pypdf import PdfReader
from docx import Document as DocxDocument

from ..config import get_llm_provider, OPENAI_API_KEY, ANTHROPIC_API_KEY, LLM_MODEL


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")


def read_resume_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return _read_pdf(file_bytes)
    if ext in {"docx", "doc"}:
        return _read_docx(file_bytes)
    # last resort - try decoding as text
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _read_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    return "\n".join(chunks).strip()


def _read_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs if p.text]
    # also pull text from tables since many resumes use them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paras.append(cell.text)
    return "\n".join(paras).strip()


def extract_fields(resume_text: str) -> dict:
    provider = get_llm_provider()
    if not provider:
        return _regex_fallback(resume_text)

    try:
        return _llm_extract(resume_text, provider)
    except Exception as e:
        # If the LLM call fails for any reason, surface partial data from regex
        # rather than failing the whole upload.
        data = _regex_fallback(resume_text)
        data["_llm_error"] = str(e)
        return data


def _llm_extract(resume_text: str, provider: str) -> dict:
    # Keep text bounded - resumes are short, but a stray giant file shouldn't blow up costs
    text = resume_text[:15000]

    system = (
        "You extract structured information from resumes. "
        "Return ONLY valid JSON with these keys: "
        "name, email, phone, company, designation, skills, confidence. "
        "`skills` is an array of strings. "
        "`company` and `designation` are the current/most recent position. "
        "`confidence` is an object mapping each of the other keys to a float in [0,1] "
        "representing how sure you are. Use null when a field is missing."
    )
    user = f"Resume text:\n\n{text}\n\nReturn the JSON object only - no prose."

    if provider == "openai":
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        model_name = LLM_MODEL or "gpt-4o-mini"
        llm = ChatOpenAI(model=model_name, api_key=OPENAI_API_KEY, temperature=0)
        resp = llm.invoke([SystemMessage(content=system), HumanMessage(content=user)])
        content = resp.content
    else:
        from langchain_anthropic import ChatAnthropic
        from langchain_core.messages import SystemMessage, HumanMessage

        model_name = LLM_MODEL or "claude-haiku-4-5-20251001"
        llm = ChatAnthropic(model=model_name, api_key=ANTHROPIC_API_KEY, temperature=0)
        resp = llm.invoke([SystemMessage(content=system), HumanMessage(content=user)])
        content = resp.content

    if isinstance(content, list):
        # Some providers return content as a list of blocks
        content = "".join(part.get("text", "") if isinstance(part, dict) else str(part) for part in content)

    return _parse_llm_json(content)


def _parse_llm_json(raw: str) -> dict:
    raw = (raw or "").strip()
    # strip code fences if the model added them
    if raw.startswith("```"):
        raw = raw.strip("`")
        # remove an optional 'json' language tag on first line
        if "\n" in raw:
            first, rest = raw.split("\n", 1)
            if first.lower().strip() in {"json", ""}:
                raw = rest
    # Some models still wrap with explanation - find the first {...} block
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1:
        return {}
    snippet = raw[start:end + 1]
    try:
        data = json.loads(snippet)
    except json.JSONDecodeError:
        return {}

    skills = data.get("skills") or []
    if isinstance(skills, str):
        skills = [s.strip() for s in re.split(r"[,;\n]", skills) if s.strip()]

    confidence = data.get("confidence") or {}
    if not isinstance(confidence, dict):
        confidence = {}

    return {
        "name": _clean(data.get("name")),
        "email": _clean(data.get("email")),
        "phone": _clean(data.get("phone")),
        "company": _clean(data.get("company")),
        "designation": _clean(data.get("designation")),
        "skills": [str(s).strip() for s in skills if s],
        "confidence": {k: float(v) for k, v in confidence.items() if _is_number(v)},
    }


def _clean(val) -> Optional[str]:
    if val is None:
        return None
    if isinstance(val, str):
        v = val.strip()
        return v or None
    return str(val)


def _is_number(v) -> bool:
    try:
        float(v)
        return True
    except (TypeError, ValueError):
        return False


def _regex_fallback(text: str) -> dict:
    """Crude offline extraction so the API still works without an LLM key."""
    if not text:
        return {"name": None, "email": None, "phone": None, "company": None,
                "designation": None, "skills": [], "confidence": {}}

    email = None
    m = EMAIL_RE.search(text)
    if m:
        email = m.group(0)

    phone = None
    m = PHONE_RE.search(text)
    if m:
        phone = re.sub(r"\s+", " ", m.group(0)).strip()

    # First non-empty line that looks like a name (no @, no digits)
    name = None
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if "@" in line or any(ch.isdigit() for ch in line):
            continue
        # name lines usually have <= 5 words
        if 1 <= len(line.split()) <= 5 and len(line) < 80:
            name = line
            break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "company": None,
        "designation": None,
        "skills": [],
        "confidence": {
            "name": 0.4 if name else 0.0,
            "email": 0.9 if email else 0.0,
            "phone": 0.7 if phone else 0.0,
        },
    }
